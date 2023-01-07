#! python

#read-docx.py

#Programa le os textos e runs do arquivo word

#TODO:

import docx

doc = docx.Document('demo.docx')
for texts in range(len(doc.paragraphs)):
  print(doc.paragraphs[texts].text + " Text")
  for run in range(len(doc.paragraphs[texts].runs)):
    print(doc.paragraphs[texts].runs[run].text + " Runs")